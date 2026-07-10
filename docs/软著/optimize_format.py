#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
优化 CanvasFlow AI Studio 软件操作说明书的排版与字体
符合软著申请规范要求
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_font(run, font_name_cn='宋体', font_name_en='Times New Roman', size=None, bold=None):
    """设置字体，中文和西文分别设置"""
    run.font.name = font_name_en
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)
    
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold

def set_paragraph_format(para, first_line_indent=True, line_spacing=1.5, alignment=None, space_before=None, space_after=None):
    """设置段落格式"""
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    
    if first_line_indent:
        # 首行缩进2字符
        pf.first_line_indent = Cm(0.74)  # 约2字符
    
    if alignment is not None:
        para.alignment = alignment
    
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)

def add_page_number(section):
    """添加页码"""
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加页码字段
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    
    set_font(run, '宋体', 'Times New Roman', size=10.5)

def optimize_document(input_path, output_path):
    """优化文档排版"""
    doc = Document(input_path)
    
    # 1. 设置页面格式
    for section in doc.sections:
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        
        # 添加页码
        add_page_number(section)
    
    # 2. 遍历所有段落，设置格式
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        text = para.text.strip()
        
        # 判断是否为标题
        is_heading = False
        heading_level = 0
        
        if style_name.startswith('Heading') or style_name.startswith('标题'):
            is_heading = True
            # 提取标题级别
            if '1' in style_name or '一' in style_name:
                heading_level = 1
            elif '2' in style_name or '二' in style_name:
                heading_level = 2
            elif '3' in style_name or '三' in style_name:
                heading_level = 3
            else:
                heading_level = 2
        
        # 根据内容判断标题
        if not is_heading and text:
            # 一级标题模式：第X章、第一篇、第二篇等
            if (text.startswith('第') and ('章' in text or '篇' in text or '部分' in text)) or \
               text.startswith('附录') or text == '目 录' or text == '目录':
                is_heading = True
                heading_level = 1
            # 二级标题模式：X.X
            elif len(text) > 2 and text[0].isdigit() and '.' in text[:5]:
                is_heading = True
                heading_level = 2
            # 三级标题模式：X.X.X
            elif text.count('.') >= 2 and text[0].isdigit():
                is_heading = True
                heading_level = 3
        
        # 设置标题格式
        if is_heading:
            if heading_level == 1:
                # 一级标题：黑体，三号，居中，段前段后间距
                for run in para.runs:
                    set_font(run, '黑体', 'Times New Roman', size=16, bold=True)
                set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                                   space_before=18, space_after=12)
            elif heading_level == 2:
                # 二级标题：黑体，小三
                for run in para.runs:
                    set_font(run, '黑体', 'Times New Roman', size=15, bold=True)
                set_paragraph_format(para, first_line_indent=False, space_before=12, space_after=6)
            else:
                # 三级标题：黑体，四号
                for run in para.runs:
                    set_font(run, '黑体', 'Times New Roman', size=14, bold=True)
                set_paragraph_format(para, first_line_indent=False, space_before=6, space_after=3)
        else:
            # 正文：宋体，小四，1.5倍行距，首行缩进
            if text:  # 非空段落
                for run in para.runs:
                    # 保留代码等特殊格式的字体
                    if not run.font.name or 'Consolas' not in run.font.name:
                        set_font(run, '宋体', 'Times New Roman', size=12)
                set_paragraph_format(para, first_line_indent=True, line_spacing=1.5)
    
    # 3. 设置表格格式
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_font(run, '宋体', 'Times New Roman', size=10.5)
                    set_paragraph_format(para, first_line_indent=False, line_spacing=1.2)
    
    # 4. 保存文档
    doc.save(output_path)
    print(f"文档已优化并保存到: {output_path}")

if __name__ == '__main__':
    input_file = r'D:\vibevideo\docs\软著\CanvasFlow-AI-Studio-软件操作说明书-V2.docx'
    output_file = r'D:\vibevideo\docs\软著\CanvasFlow-AI-Studio-软件操作说明书-V2-排版优化版.docx'
    
    optimize_document(input_file, output_file)
